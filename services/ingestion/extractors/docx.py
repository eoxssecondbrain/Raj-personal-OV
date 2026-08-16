from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from . import ExtractionError


def extract(file_path):
    try:
        doc = Document(file_path)
    except PackageNotFoundError as e:
        raise ExtractionError(f"corrupted or not a valid docx: {e}") from e

    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    if not text.strip():
        raise ExtractionError("no extractable text found in document")
    return text
